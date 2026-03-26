import os
import json
import sqlite3
from datetime import datetime, date
from pathlib import Path

import streamlit as st
from docx import Document

st.set_page_config(page_title="신입강사 멘토링 관리", layout="wide")

BASE_DIR = Path(".")
DATA_DIR = BASE_DIR / "data"
DOCS_DIR = BASE_DIR / "generated_docs"
DB_PATH = DATA_DIR / "mentoring.db"

DATA_DIR.mkdir(exist_ok=True)
DOCS_DIR.mkdir(exist_ok=True)


def now_str():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def get_conn():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_conn()
    cur = conn.cursor()

    cur.execute("""
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        login_id TEXT UNIQUE NOT NULL,
        name TEXT NOT NULL,
        role TEXT NOT NULL,
        password TEXT DEFAULT '0000',
        created_at TEXT NOT NULL
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS assignments (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        mentor_user_id INTEGER NOT NULL,
        mentee_user_id INTEGER NOT NULL UNIQUE,
        assigned_at TEXT NOT NULL
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        related_mentee_user_id INTEGER,
        doc_type TEXT NOT NULL,
        version_label TEXT,
        content_json TEXT NOT NULL,
        file_path TEXT,
        submitted_at TEXT NOT NULL
    )
    """)

    conn.commit()

    cur.execute("SELECT COUNT(*) AS cnt FROM users")
    if cur.fetchone()["cnt"] == 0:
        users = [
            ("admin001", "관리자", "admin", "0000", now_str()),
            ("김소영101", "김소영", "mentee", "0000", now_str()),
            ("박민지102", "박민지", "mentee", "0000", now_str()),
            ("이준호201", "이준호", "mentor", "0000", now_str()),
            ("최나영202", "최나영", "mentor", "0000", now_str()),
        ]
        cur.executemany("""
        INSERT INTO users (login_id, name, role, password, created_at)
        VALUES (?, ?, ?, ?, ?)
        """, users)
        conn.commit()

    conn.close()


def authenticate(login_id, password, role):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("""
    SELECT * FROM users WHERE login_id = ? AND password = ? AND role = ?
    """, (login_id, password, role))
    row = cur.fetchone()
    conn.close()
    return row


def list_users_by_role(role):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT * FROM users WHERE role = ? ORDER BY name", (role,))
    rows = cur.fetchall()
    conn.close()
    return rows


def get_assignment_for_mentee(mentee_user_id):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("""
    SELECT a.*, u.name AS mentor_name, u.login_id AS mentor_login_id
    FROM assignments a
    JOIN users u ON a.mentor_user_id = u.id
    WHERE a.mentee_user_id = ?
    """, (mentee_user_id,))
    row = cur.fetchone()
    conn.close()
    return row


def get_assigned_mentees_for_mentor(mentor_user_id):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("""
    SELECT a.*, u.name AS mentee_name, u.login_id AS mentee_login_id
    FROM assignments a
    JOIN users u ON a.mentee_user_id = u.id
    WHERE a.mentor_user_id = ?
    ORDER BY u.name
    """, (mentor_user_id,))
    rows = cur.fetchall()
    conn.close()
    return rows


def upsert_assignment(mentor_user_id, mentee_user_id):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT id FROM assignments WHERE mentee_user_id = ?", (mentee_user_id,))
    row = cur.fetchone()
    if row:
        cur.execute("""
        UPDATE assignments
        SET mentor_user_id = ?, assigned_at = ?
        WHERE mentee_user_id = ?
        """, (mentor_user_id, now_str(), mentee_user_id))
    else:
        cur.execute("""
        INSERT INTO assignments (mentor_user_id, mentee_user_id, assigned_at)
        VALUES (?, ?, ?)
        """, (mentor_user_id, mentee_user_id, now_str()))
    conn.commit()
    conn.close()


def ensure_user_dir(login_id):
    p = DOCS_DIR / login_id
    p.mkdir(parents=True, exist_ok=True)
    return p


def save_docx_from_content(doc_type, content, file_path):
    doc = Document()
    doc.add_heading(doc_type, level=1)

    for k, v in content.items():
        if isinstance(v, list):
            doc.add_paragraph(f"{k}: {', '.join([str(x) for x in v])}")
        elif isinstance(v, dict):
            doc.add_paragraph(f"[{k}]")
            for kk, vv in v.items():
                doc.add_paragraph(f"{kk}: {vv}")
        else:
            doc.add_paragraph(f"{k}: {v}")

    doc.save(file_path)


def save_document_record(user, doc_type, content, related_mentee_user_id=None, version_label=None):
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    user_dir = ensure_user_dir(user["login_id"])
    filename = f"{doc_type}_{version_label or '기본'}_{ts}.docx".replace("/", "_")
    file_path = user_dir / filename

    save_docx_from_content(doc_type, content, file_path)

    json_path = user_dir / f"{doc_type}_{version_label or '기본'}_{ts}.json".replace("/", "_")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(content, f, ensure_ascii=False, indent=2)

    conn = get_conn()
    cur = conn.cursor()
    cur.execute("""
    INSERT INTO documents
    (user_id, related_mentee_user_id, doc_type, version_label, content_json, file_path, submitted_at)
    VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (
        user["id"],
        related_mentee_user_id,
        doc_type,
        version_label,
        json.dumps(content, ensure_ascii=False),
        str(file_path),
        now_str()
    ))
    conn.commit()
    conn.close()
    return str(file_path)


def list_documents_for_user(user_id):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT * FROM documents WHERE user_id = ? ORDER BY submitted_at DESC", (user_id,))
    rows = cur.fetchall()
    conn.close()
    return rows


def list_all_documents():
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("""
    SELECT d.*, u.name AS writer_name, u.role AS writer_role
    FROM documents d
    JOIN users u ON d.user_id = u.id
    ORDER BY d.submitted_at DESC
    """)
    rows = cur.fetchall()
    conn.close()
    return rows


def render_login():
    st.title("신입강사 멘토링 관리 시스템")
    role_label = st.selectbox("역할", ["멘티", "멘토", "관리자"])
    role_map = {"멘티": "mentee", "멘토": "mentor", "관리자": "admin"}

    login_id = st.text_input("아이디")
    password = st.text_input("비밀번호", type="password", value="0000")

    if st.button("로그인"):
        user = authenticate(login_id, password, role_map[role_label])
        if user:
            st.session_state["user"] = dict(user)
            st.rerun()
        else:
            st.error("로그인 정보가 올바르지 않습니다.")


def render_sidebar():
    user = st.session_state["user"]
    with st.sidebar:
        st.write(f"**이름:** {user['name']}")
        st.write(f"**아이디:** {user['login_id']}")
        st.write(f"**권한:** {user['role']}")
        if st.button("로그아웃"):
            st.session_state.clear()
            st.rerun()


def form_mentee_agreement():
    st.subheader("신입강사 멘토링 참여 서약서")

    st.markdown("""
    ※ 본 멘토링 프로그램은 신입강사의 초기 적응 지원과 수업 역량 강화를 목적으로 하며,
    멘토의 지도와 점검을 통해 안정적인 정착을 돕기 위한 과정입니다.

    1. 본인은 배정담당자의 추천 또는 회사의 운영 기준에 따라 본 멘토링 프로그램 대상자로 선정되었음을 숙지하였습니다.
    2. 본인은 본 멘토링 프로그램이 본인의 자발적인 참여 의사에 따라 진행됨을 확인하였습니다.
    3. 본인은 멘티로서 멘토 및 교관의 지도와 안내 사항을 성실히 이행하겠습니다.
    4. 본인은 멘토링 진행 과정에서 요청되는 수업계획서 작성, 수업 운영안 제출, 피드백 반영 등 제반 사항에 성실히 참여하겠습니다.
    5. 본인은 멘토의 정당한 지도 및 요청 사항을 반복적으로 이행하지 않을 경우, 멘토링이 조기 종료될 수 있음을 확인하였습니다.
    6. 본인은 특별한 사유 없이 멘토링을 중도 포기하거나 운영 기준을 성실히 따르지 않을 경우, 향후 배정 및 운영상 불이익이 발생할 수 있음을 확인하였습니다.
    7. 본인은 첫 수업 및 멘토링 진행 과정에서 확인된 보완 사항에 대해 성실히 개선하고, 필요한 경우 추가 지도 및 보완 절차에 협조하겠습니다.
    8. 본인은 멘토링 과정을 충실히 이행한 경우, 회사 운영 기준에 따라 우선 배정 검토, 추가 성장 기회 제공 등 후속 지원이 이루어질 수 있음을 확인하였습니다.

    상기 본인은 이지수능교육 강사로서 위 내용을 충분히 숙지하고 이해하였으며,
    이를 성실히 준수할 것을 서약합니다.
    """)

    st.divider()

    agree = st.checkbox("위 내용을 모두 확인하였으며 이에 동의합니다.", key="mentee_agree_check")
    signed_date = st.date_input("작성일", value=date.today(), key="mentee_agree_date")
    name = st.text_input("성명", key="mentee_agree_name")    

    return {
        "서약서 동의 여부": agree,
        "작성일": str(signed_date),
        "성명": name
    }


def form_mentee_info():
    st.subheader("신입강사용 사전 정보지")
    data = {
        "성명": st.text_input("성명"),
        "별칭": st.text_input("별칭"),
        "출생년도": st.text_input("출생년도"),
        "입사년도": st.text_input("입사년도"),
        "연락처": st.text_input("연락처"),
        "거주지": st.text_input("거주지"),
        "주요 배정 가능 지역": st.text_input("주요 배정 가능 지역"),
        "최종학력 / 전공": st.text_input("최종학력 / 전공"),
        "전임 여부": st.radio("전임 여부", ["전임", "비전임"], horizontal=True),
        "강의 가능 과목": st.text_input("강의 가능 과목"),
        "강의 가능 학년": st.text_input("강의 가능 학년"),
        "총 강의 경력": st.text_input("총 강의 경력"),
        "이전 근무 기관 또는 강의 경험": st.text_area("이전 근무 기관 또는 강의 경험"),
        "주요 강의 과목": st.text_input("주요 강의 과목"),
        "주요 강의 대상": st.text_input("주요 강의 대상"),
        "가장 자신 있는 수업 형태": st.text_area("가장 자신 있는 수업 형태"),
        "가장 어려움을 느끼는 수업 형태": st.text_area("가장 어려움을 느끼는 수업 형태"),
        "본인이 생각하는 강점": st.text_area("본인이 생각하는 강점"),
        "학생과의 소통에서 자신 있는 점": st.text_area("학생과의 소통에서 자신 있는 점"),
        "학부모와의 소통에서 자신 있는 점": st.text_area("학부모와의 소통에서 자신 있는 점"),
        "수업 운영 시 보완 필요 점": st.text_area("수업 운영 시 보완 필요 점"),
        "첫 수업에서 가장 걱정되는 부분": st.text_area("첫 수업에서 가장 걱정되는 부분"),
        "현재 배정 여부": st.radio("현재 배정 여부", ["배정 완료", "배정 대기"], horizontal=True),
        "첫 수업 예정일": str(st.date_input("첫 수업 예정일", value=date.today())),
        "배정 과목": st.text_input("배정 과목"),
        "학생 학년": st.text_input("학생 학년"),
        "학생 특이사항": st.text_area("학생 특이사항"),
        "학부모 요청사항": st.text_area("학부모 요청사항", key="parent_request_1"),
        "배정담당자 전달사항": st.text_area("배정담당자 전달사항"),
        "도움받고 싶은 부분": st.multiselect(
            "도움받고 싶은 부분",
            ["첫 수업 준비", "수업계획서 작성", "학생 수준 진단", "학부모 상담 및 소통",
             "과제 및 피드백 운영", "설명력 및 시간 배분", "중단 예방 및 유지 관리", "기타"]
        ),
        "멘토에게 특별히 확인받고 싶은 사항": st.text_area("멘토에게 특별히 확인받고 싶은 사항"),
        "멘토링 진행 시 요청사항": st.text_area("멘토링 진행 시 요청사항"),
        "과거 수업 중 가장 잘 진행되었던 사례": st.text_area("과거 수업 중 가장 잘 진행되었던 사례"),
        "과거 수업 중 어려움을 겪었던 사례": st.text_area("과거 수업 중 어려움을 겪었던 사례"),
        "인계 또는 중단 경험 유무": st.radio("인계 또는 중단 경험 유무", ["있음", "없음"], horizontal=True),
        "인계/중단 주요 사유": st.text_area("인계/중단 주요 사유"),
        "당시 보완 필요 사항": st.text_area("당시 보완 필요 사항"),
        "멘토에게 미리 공유하고 싶은 내용": st.text_area("멘토에게 미리 공유하고 싶은 내용"),
        "현재 가장 필요한 지원": st.text_area("현재 가장 필요한 지원"),
        "기타 참고 사항": st.text_area("기타 참고 사항"),
        "작성 확인 동의": st.checkbox("사실에 근거하여 작성하였으며 기초 자료 활용에 동의합니다."),
        "작성일": str(st.date_input("작성일", value=date.today(), key="mentee_info_date"))
        }

    st.markdown("### 자기 점검 체크")
    self_checks = {}
    items = [
        "1~8회차 수업 흐름을 설계할 수 있다.",
        "첫 수업 90분 구조를 구성할 수 있다.",
        "학부모 OT 또는 사전 통화를 진행할 수 있다.",
        "첫 수업 후 학부모 피드백을 전달할 수 있다.",
        "학생 수준에 맞게 수업 난이도를 조절할 수 있다.",
        "과제 및 피드백을 안정적으로 운영할 수 있다.",
        "학생·학부모 만족도를 관리하며 수업을 유지할 수 있다.",
        "중단 위험 신호를 인지하고 대응할 수 있다."
    ]
    for i, item in enumerate(items):
        self_checks[item] = st.radio(item, ["그렇다", "보통이다", "어렵다"], horizontal=True, key=f"self_{i}")
    data["자기 점검"] = self_checks
    return data


def form_lesson_plan_basic():
    st.subheader("수업계획서")
    return {
        "학생명": st.text_input("학생명"),
        "학교 / 학년": st.text_input("학교 / 학년"),
        "과목": st.text_input("과목"),
        "담당강사": st.text_input("담당강사"),
        "연락처(HP)": st.text_input("연락처(HP)"),
        "작성일": str(st.date_input("작성일", value=date.today(), key="plan_basic_date")),
        "현재 성적 및 수준": st.text_area("현재 성적 및 수준"),
        "학습 성향": st.text_area("학습 성향"),
        "강점": st.text_area("강점"),
        "보완 필요 영역": st.text_area("보완 필요 영역"),
        "학부모 요청사항": st.text_area("학부모 요청사항"),
        "단기 목표(1~2개월)": st.text_area("단기 목표(1~2개월)"),
        "중기 목표(3~6개월)": st.text_area("중기 목표(3~6개월)"),
        "수업 운영 방향": st.text_area("수업 운영 방향"),
        "기본 교재": st.text_input("기본 교재"),
        "부교재 / 프린트": st.text_input("부교재 / 프린트"),
        "과제 운영 방식": st.text_area("과제 운영 방식"),
    }

def form_first_class_plan():
    st.subheader("첫 수업 운영 계획")
    return {
        "사전 통화 여부 / OT 진행 내용": st.text_area("사전 통화 여부 / OT 진행 내용"),
        "학생 진단 방식": st.text_area("학생 진단 방식"),
        "첫 수업 90분 구성": st.text_area("첫 수업 90분 구성"),
        "첫 수업 후 학부모 전달 내용": st.text_area("첫 수업 후 학부모 전달 내용"),
        "2회차 전 보완 예정 사항": st.text_area("2회차 전 보완 예정 사항"),
        "커리큘럼 적정성": st.text_area("커리큘럼 적정성"),
        "교재 및 과제 수준": st.text_area("교재 및 과제 수준"),
        "첫 수업 운영 구조": st.text_area("첫 수업 운영 구조"),
        "학부모 소통 문구 / 멘토 의견": st.text_area("학부모 소통 문구 / 멘토 의견"),
        "작성일": str(st.date_input("작성일", value=date.today(), key="first_class_date")),
    }

    st.markdown("### 회차별 수업 계획")
    sessions = {}
    for i in range(1, 9):
        with st.expander(f"{i}회차"):
            sessions[f"{i}회차"] = {
                "강의 주제": st.text_input("강의 주제", key=f"topic_{i}"),
                "학습 내용": st.text_area("학습 내용", key=f"content_{i}"),
                "회차 목표": st.text_area("회차 목표", key=f"goal_{i}"),
                "과제": st.text_area("과제", key=f"hw_{i}"),
                "학부모 피드백 포인트": st.text_area("학부모 피드백 포인트", key=f"feedback_{i}")
            }
    data["회차별 수업 계획"] = sessions
    return data

def form_session_plan():
    st.subheader("회차별 수업계획")

    session_no = st.selectbox("회차 선택", list(range(1, 9)), key="session_plan_no")

    session_data = {
        "회차": session_no,
        "강의 주제": st.text_input("강의 주제", key=f"topic_{session_no}"),
        "학습 내용": st.text_area("학습 내용", key=f"content_{session_no}"),
        "회차 목표": st.text_area("회차 목표", key=f"goal_{session_no}"),
        "과제": st.text_area("과제", key=f"hw_{session_no}"),
        "학부모 피드백 포인트": st.text_area("학부모 피드백 포인트", key=f"feedback_{session_no}"),
        "작성일": str(st.date_input("작성일", value=date.today(), key=f"session_date_{session_no}"))
    }

    return session_data


def form_mentor_info():
    st.subheader("멘토작성용 사전 정보지")

    return {
        # 1. 멘티 기본 정보
        "멘티명": st.text_input("멘티명"),
        "별칭": st.text_input("별칭"),
        "과목": st.text_input("과목"),
        "입사년도": st.text_input("입사년도"),
        "거주지 / 주요 배정지역": st.text_input("거주지 / 주요 배정지역"),
        "강의 가능 학년": st.text_input("강의 가능 학년"),
        "멘토명": st.text_input("멘토명"),
        "교관명": st.text_input("교관명"),

        # 2. 배정 및 멘토링 정보
        "멘토링 시작일": str(st.date_input("멘토링 시작일", value=date.today(), key="mentor_info_start")),
        "배정 과목": st.text_input("배정 과목"),
        "첫 수업 예정일": str(st.date_input("첫 수업 예정일", value=date.today(), key="mentor_info_first_class")),
        "학생명 / 학년": st.text_input("학생명 / 학년"),
        "멘토링 구분": st.radio(
            "멘토링 구분",
            ["신규 배정", "재배정", "재입사", "기타"],
            horizontal=True
        ),
        "중점 관리 항목": st.multiselect(
            "중점 관리 항목",
            [
                "첫 수업 준비",
                "학부모 소통",
                "수업계획서 작성",
                "과제 및 피드백 운영",
                "설명력 및 시간 배분",
                "중단 위험 관리",
                "기타"
            ]
        ),
        "멘토링 필요 사유": st.text_area("멘토링 필요 사유"),

        # 3. 강의 및 배정 이력 요약
        "현재 진행 수업 / 배정 현황 요약": st.text_area("현재 진행 수업 / 배정 현황 요약"),
        "과거 인계 이력 요약": st.text_area("과거 인계 이력 요약"),
        "과거 중단 이력 요약": st.text_area("과거 중단 이력 요약"),
        "배정담당자 전달사항 및 유의사항": st.text_area("배정담당자 전달사항 및 유의사항"),

        # 4. 실적 및 운영 참고 정보
        "최근 연간 실적 요약": st.text_area("최근 연간 실적 요약"),
        "전임 구분": st.text_input("전임 구분"),
        "학생·학부모 소통 강점": st.text_area("학생·학부모 소통 강점"),
        "수업 운영 강점": st.text_area("수업 운영 강점"),
        "보완 필요 사항": st.text_area("보완 필요 사항"),
        "주의 필요 이력": st.text_area("주의 필요 이력"),
        "기타 참고사항": st.text_area("기타 참고사항"),

        # 5. 첫 수업 사전 점검
        "수업계획서 제출 여부": st.radio(
            "수업계획서 제출 여부",
            ["완료", "확인 필요", "미완료"],
            horizontal=True
        ),
        "첫 수업 운영안 제출 여부": st.radio(
            "첫 수업 운영안 제출 여부",
            ["완료", "확인 필요", "미완료"],
            horizontal=True
        ),
        "OT/사전 통화 멘트 준비 여부": st.radio(
            "OT/사전 통화 멘트 준비 여부",
            ["완료", "확인 필요", "미완료"],
            horizontal=True
        ),
        "학생 수준 파악 여부": st.radio(
            "학생 수준 파악 여부",
            ["완료", "확인 필요", "미완료"],
            horizontal=True
        ),
        "학부모 요청사항 확인 여부": st.radio(
            "학부모 요청사항 확인 여부",
            ["완료", "확인 필요", "미완료"],
            horizontal=True
        ),
        "멘토 사전 코멘트": st.text_area("멘토 사전 코멘트"),

        # 6. 초기 리스크 및 관리 계획
        "초기 적응 예상도": st.radio(
            "초기 적응 예상도",
            ["높음", "보통", "낮음"],
            horizontal=True
        ),
        "중단 위험도": st.radio(
            "중단 위험도",
            ["낮음", "보통", "높음"],
            horizontal=True
        ),
        "초기 중점 관리 방향": st.text_area("초기 중점 관리 방향"),
        "첫 수업 전 반드시 보완이 필요한 사항": st.text_area("첫 수업 전 반드시 보완이 필요한 사항"),
        "첫 수업 후 확인 예정 포인트": st.text_area("첫 수업 후 확인 예정 포인트"),

        # 7. 멘토 종합 의견
        "멘토링 시 유의사항": st.text_area("멘토링 시 유의사항"),
        "기타 의견": st.text_area("기타 의견"),

        # 작성 정보
        "작성일": str(st.date_input("작성일", value=date.today(), key="mentor_info_date")),
        "멘토 성명": st.text_input("멘토 성명"),
        }


def form_mentor_report():
    st.subheader("멘토링 보고서")

    return {
        # 1. 기본 정보
        "멘토링 구간(Stage)": st.radio(
            "멘토링 구간(Stage)",
            ["Stage 1", "Stage 2", "Stage 3"],
            horizontal=True
        ),
        "관리 회차": st.text_input("관리 회차"),
        "상담 방식": st.multiselect(
            "상담 방식",
            ["전화", "대면", "메신저 병행"]
        ),
        "일시": str(st.date_input("일시", value=date.today(), key="mentor_report_date")),
        "소요시간": st.text_input("소요시간"),
        "멘토링 장소": st.text_input("멘토링 장소"),
        "중단 위험 여부": st.radio(
            "중단 위험 여부",
            ["낮음", "보통", "높음"],
            horizontal=True
        ),
        "상담 결과": st.radio(
            "상담 결과",
            ["유지", "보완", "인계/중단"],
            horizontal=True
        ),

        # 2. 멘토·멘티 정보
        "멘토 성명": st.text_input("멘토 성명"),
        "멘토 소속/과목": st.text_input("멘토 소속/과목"),
        "멘티 성명": st.text_input("멘티 성명"),
        "멘티 소속/과목": st.text_input("멘티 소속/과목"),
        "학생명": st.text_input("학생명"),
        "수업 과목/학년": st.text_input("수업 과목/학년"),

        # 3. 멘토링 목적
        "멘토링 목적": st.text_area("멘토링 목적"),

        # 4. 멘토링 상세 내용
        "수업 준비도": st.text_area("① 수업 준비도"),
        "수업 진행 능력": st.text_area("② 수업 진행 능력"),
        "학생·학부모 소통": st.text_area("③ 학생·학부모 소통"),
        "확인된 문제점": st.text_area("④ 확인된 문제점"),
        "멘토의 지도·조언": st.text_area("⑤ 멘토의 지도·조언"),

        # 5. 멘토 평가 및 종합 의견
        "변화 및 개선 여부": st.text_area("① 변화 및 개선 여부"),
        "멘티 태도 및 성실도": st.text_area("② 멘티 태도 및 성실도"),
        "향후 지도 필요 사항": st.text_area("③ 향후 지도 필요 사항"),
        "종합 의견": st.text_area("④ 종합 의견"),

        # 6. 멘토링 결과 증빙
        "첨부 내역": st.multiselect(
            "첨부 내역",
            ["카카오톡 대화 캡처", "통화 기록", "상담 메모", "녹취 파일", "사진", "수업 자료", "피드백 자료"]
        ),
        "첨부 자료 설명 1": st.text_area("첨부 자료 설명 1"),
        "첨부 자료 설명 2": st.text_area("첨부 자료 설명 2"),
        "첨부 자료 설명 3": st.text_area("첨부 자료 설명 3"),

        # 작성 정보
        "작성일": str(st.date_input("작성일", value=date.today(), key="mentor_report_written_date")),
    }


def form_mentor_checklist():
    st.subheader("회차별 멘토 확인표")

    data = {
        "멘티명": st.text_input("멘티명"),
        "과목": st.text_input("과목"),
        "멘토명": st.text_input("멘토명"),
        "시작일": str(st.date_input("시작일", value=date.today(), key="mentor_check_start_date")),
        "학생명": st.text_input("학생명"),
    }

    stage_sections = {
        "Stage 1": {
            "첫 수업 전 확인 사항": [
                "1~3회차 수업계획서 작성 여부",
                "첫 수업 90분 구성안 작성 여부",
                "학생 수준 및 학부모 요청사항 파악 여부",
                "OT/사전 통화 멘트 준비 여부",
                "교재 및 과제 기준 설정 여부",
                "첫 수업 후 학부모 피드백 계획 수립 여부",
            ],
            "첫 수업 후 확인 사항": [
                "실제 수업 진행 흐름 확인",
                "학생 반응 및 참여도 확인",
                "학부모 반응 및 초기 만족도 확인",
                "담임 모니터링 필요 여부 확인",
                "수업계획서와 실제 수업의 차이 확인",
            ],
        },
        "Stage 2": {
            "3회차 후 확인 사항": [
                "4~5회차 진행 후 수업 안정화 여부",
                "학생 반응 및 출결·과제 수행 상태 확인",
                "학부모 신뢰 형성 여부",
                "수업 중단 위험 요소 확인",
                "설명 방식, 시간 배분, 과제 운영 보완 필요 여부",
            ],
            "4~5회차 확인 사항": [
                "1개월 운영 결과 회고",
                "설명 방식 및 시간 배분 보정 여부",
                "과제·피드백 루틴 안정화 여부",
                "중간 상담 준비 상태",
                "중단 위험 신호 확인",
            ],
        },
        "Stage 3": {
            "6~8회차 확인 사항": [
                "2개월 운영 성과 정리 여부",
                "학생·학부모 만족도 추이 확인",
                "장기 유지 가능성 점검",
                "추가 타임 또는 과목 확장 가능성 검토",
                "3~6개월 운영 방향 설계 여부",
            ],
        },
    }

    detail = {}

    for stage_name, sections in stage_sections.items():
        st.markdown(f"## {stage_name}")
        detail[stage_name] = {}

        for section_name, items in sections.items():
            st.markdown(f"### {section_name}")
            detail[stage_name][section_name] = {}

            for idx, item in enumerate(items):
                with st.expander(item):
                    detail[stage_name][section_name][item] = {
                        "확인 결과": st.text_area(
                            f"{item} - 확인 결과",
                            key=f"{stage_name}_{section_name}_{idx}_result"
                        ),
                        "보완 사항": st.text_area(
                            f"{item} - 보완 사항",
                            key=f"{stage_name}_{section_name}_{idx}_action"
                        ),
                        "확인일": str(
                            st.date_input(
                                f"{item} - 확인일",
                                value=date.today(),
                                key=f"{stage_name}_{section_name}_{idx}_date"
                            )
                        ),
                    }

    data["확인표 상세"] = detail

    st.markdown("## 첫 수업 후 분기별 판단")
    data["첫 수업 후 분기별 판단"] = st.radio(
        "판단 결과",
        ["만족", "보완 필요", "불만족 또는 인계 발생"],
        horizontal=True,
        key="mentor_check_branching"
    )

    data["특이사항"] = st.text_area("특이사항")
    data["후속 조치"] = st.text_area("후속 조치")
    data["멘토 종합 의견"] = st.text_area("멘토 종합 의견")

    return data


def form_mentor_eval():
    st.subheader("신입강사 멘토링 평가표")
    data = {
        "멘티 이름": st.text_input("멘티 이름"),
        "과목": st.text_input("과목"),
        "멘토 이름": st.text_input("멘토 이름"),
        "평가자": st.text_input("평가자"),
        "멘토링 기간": st.text_input("멘토링 기간"),
        "비고": st.text_area("비고")
    }

    items = [
        ("Stage 1", "수업 준비", "1~8회차 커리큘럼을 설계하였는가"),
        ("Stage 1", "수업 준비", "첫 수업 90분 구조를 설계하였는가"),
        ("Stage 1", "수업 준비", "교재 및 과제 기준을 설정하였는가"),
        ("Stage 1", "학부모 소통", "OT/사전 통화 멘트를 준비하였는가"),
        ("Stage 1", "첫 수업 운영", "첫 수업 후 피드백 계획을 수립하였는가"),
        ("Stage 2", "수업 운영", "첫 수업 결과를 반영하여 계획서를 수정하였는가"),
        ("Stage 2", "수업 운영", "학생 수준에 맞춰 난이도를 조정하였는가"),
        ("Stage 2", "수업 운영", "설명 방식과 시간 배분을 보완하였는가"),
        ("Stage 2", "과제 관리", "과제·피드백 루틴이 안정적으로 운영되는가"),
        ("Stage 2", "상담 대응", "중간 상담 준비가 되어 있는가"),
        ("Stage 3", "유지 관리", "2개월 운영 성과를 정리할 수 있는가"),
        ("Stage 3", "유지 관리", "학생·학부모 만족도를 안정적으로 관리하는가"),
        ("Stage 3", "확장 설계", "3~6개월 장기 로드맵을 제시할 수 있는가"),
        ("Stage 3", "확장 설계", "추가 타임/과목 확장 가능성을 검토할 수 있는가"),
        ("종합", "종합 평가", "안정적인 수업 유지가 가능한 수준에 도달하였는가"),
    ]

    scores = {}
    for i, (stage, area, item) in enumerate(items, start=1):
        scores[f"{i}. {item}"] = {
            "평가 시점": stage,
            "평가영역": area,
            "점수": st.radio(f"{i}. {item}", ["A", "B", "C"], horizontal=True, key=f"eval_{i}")
        }
    data["평가 항목"] = scores

    data["강점"] = st.text_area("강점")
    data["보완 필요 사항"] = st.text_area("보완 필요 사항")
    data["우선 지도 항목"] = st.text_area("우선 지도 항목")
    data["후속 조치"] = st.multiselect(
        "후속 조치",
        ["추가 멘토링 필요", "다음 Stage 진행 가능", "집중 관리 필요", "운영 협의 필요"]
    )
    data["최종 판정"] = st.radio(
        "최종 판정",
        ["안정적 운영 가능", "보완 후 운영 가능", "추가 점검 필요"],
        horizontal=True
    )
    return data


def render_mentee():
    user = st.session_state["user"]
    st.title(f"{user['name']} 님, 환영합니다.")

    assignment = get_assignment_for_mentee(user["id"])
    if assignment:
        st.success(f"배정된 멘토: {assignment['mentor_name']}")
    else:
        st.info("아직 멘토가 배정되지 않았습니다.")

    tab1, tab2, tab3 = st.tabs(["서류 작성", "수업 계획서", "제출 내역"])

    with tab1:
        doc = st.selectbox(
            "문서 선택",
            ["멘토링 참여서약서", "사전 정보지"],
            key="mentee_tab1_doc"
        )

        content = None
        version_label = None

        if doc == "멘토링 참여서약서":
            content = form_mentee_agreement()
        elif doc == "사전 정보지":
            content = form_mentee_info()
            version_label = "초안"

        if st.button("저장", key="mentee_tab1_save", use_container_width=True):
            if doc == "멘토링 참여서약서" and not content.get("서약서 동의 여부", False):
                st.warning("서약 내용 동의가 필요합니다.")
            else:
                path = save_document_record(
                    user=user,
                    doc_type=doc,
                    content=content,
                    related_mentee_user_id=user["id"],
                    version_label=version_label
                )
                st.success(f"저장 완료: {path}")

    with tab2:
        doc = st.selectbox(
            "작성 문서 선택",
            ["수업계획서", "첫 수업 운영 계획", "회차별 수업계획"],
            key="after_first_class_doc"
        )

        revised = None
        save_doc_type = None
        save_version = None

        if doc == "수업계획서":
            revised = form_lesson_plan_basic()
            save_doc_type = "수업계획서"
            save_version = "수정본"
        elif doc == "첫 수업 운영 계획":
            revised = form_first_class_plan()
            save_doc_type = "첫 수업 운영 계획"
            save_version = "수정본"
        elif doc == "회차별 수업계획":
            revised = form_session_plan()
            save_doc_type = "회차별 수업계획"
            save_version = f"{revised['회차']}회차"

        if st.button("수정본 저장", key="save_revised", use_container_width=True):
            path = save_document_record(
                user=user,
                doc_type=save_doc_type,
                content=revised,
                related_mentee_user_id=user["id"],
                version_label=save_version
            )
            st.success(f"수정본 저장 완료: {path}")

    with tab3:
        docs = list_documents_for_user(user["id"])
        if not docs:
            st.info("제출한 문서가 없습니다.")
        else:
            for d in docs:
                st.write(f"- {d['doc_type']} / {d['version_label'] or '-'} / {d['submitted_at']}")


def render_mentor():
    user = st.session_state["user"]
    st.title(f"{user['name']} 님, 환영합니다.")
    assigned = get_assigned_mentees_for_mentor(user["id"])

    tab1, tab2, tab3 = st.tabs(["문서 작성", "배정 멘티", "내 문서"])

    with tab1:
        doc = st.selectbox("작성 문서 선택", [
            "멘토작성용 사전 정보지",
            "멘토링 보고서",
            "회차별 멘토 확인표",
            "신입강사 멘토링 평가표"
        ])

        related_mentee_id = None
        if assigned:
            options = {f"{m['mentee_name']} ({m['mentee_login_id']})": m["mentee_user_id"] for m in assigned}
            selected = st.selectbox("관련 멘티 선택", list(options.keys()))
            related_mentee_id = options[selected]

        if doc == "멘토작성용 사전 정보지":
            content = form_mentor_info()
        elif doc == "멘토링 보고서":
            content = form_mentor_report()
        elif doc == "회차별 멘토 확인표":
            content = form_mentor_checklist()
        else:
            content = form_mentor_eval()

        if st.button("저장", key="mentor_save", use_container_width=True):
            path = save_document_record(
                user=user,
                doc_type=doc,
                content=content,
                related_mentee_user_id=related_mentee_id
            )
            st.success(f"저장 완료: {path}")

    with tab2:
        if not assigned:
            st.info("배정된 멘티가 없습니다.")
        else:
            for m in assigned:
                st.write(f"- {m['mentee_name']}")

    with tab3:
        docs = list_documents_for_user(user["id"])
        for d in docs:
            st.write(f"- {d['doc_type']} / {d['submitted_at']}")


def render_admin():
    st.title("관리자 화면")
    mentees = list_users_by_role("mentee")
    mentors = list_users_by_role("mentor")

    tab1, tab2 = st.tabs(["배정 관리", "전체 문서 현황"])

    with tab1:
        mentee_map = {f"{m['name']} ({m['login_id']})": m["id"] for m in mentees}
        mentor_map = {f"{m['name']} ({m['login_id']})": m["id"] for m in mentors}

        selected_mentee = st.selectbox("멘티 선택", list(mentee_map.keys()))
        selected_mentor = st.selectbox("멘토 선택", list(mentor_map.keys()))

        if st.button("배정 저장", use_container_width=True):
            upsert_assignment(mentor_map[selected_mentor], mentee_map[selected_mentee])
            st.success("배정 완료")
            st.rerun()

    with tab2:
        docs = list_all_documents()
        for d in docs:
            st.write(f"- {d['writer_name']} / {d['writer_role']} / {d['doc_type']} / {d['version_label'] or '-'} / {d['submitted_at']}")


def main():
    init_db()

    if "user" not in st.session_state:
        render_login()
        return

    render_sidebar()
    role = st.session_state["user"]["role"]
    if role == "mentee":
        render_mentee()
    elif role == "mentor":
        render_mentor()
    else:
        render_admin()


if __name__ == "__main__":
    main()